"""
企业信息处理模块
支持文档提取、文本解析、对话引导三种信息收集方式
"""

import os
import re
import json
import asyncio
from typing import Dict, Any, List, Optional, Tuple
import tempfile

from common.log import logger


class InfoProcessor:
    """企业信息处理器"""
    
    REQUIRED_FIELDS = [
        "company_name",          # 企业名称
        "establishment_year",    # 成立年份
        "industry",              # 所属行业
        "main_business",         # 主营业务
        "revenue_2023",          # 2023年营收
        "main_revenue_2023",     # 2023年主营业务收入
        "rd_expense_2023",       # 2023年研发费用
        "employee_count",        # 员工总数
        "rd_employee_count",     # 研发人员数
        "patent_count",          # 专利数量
        "main_products",         # 主要产品
        "market_share",          # 市场占有率
    ]
    
    OPTIONAL_FIELDS = [
        "revenue_2021", "revenue_2022",
        "main_revenue_2021", "main_revenue_2022",
        "rd_expense_2021", "rd_expense_2022",
        "asset_total", "liability_total",
        "core_customers", "certifications",
        "rd_institution", "management_system"
    ]
    
    def __init__(self):
        self.conversation_steps = [
            ("company_name", "请问企业全称是什么？"),
            ("establishment_year", "企业成立时间是哪一年？"),
            ("industry", "企业所属行业是什么？（如：智能制造、生物医药等）"),
            ("main_business", "主营业务具体是什么？"),
            ("revenue_2023", "2023年营业收入是多少万元？"),
            ("main_revenue_2023", "2023年主营业务收入是多少万元？"),
            ("rd_expense_2023", "2023年研发费用是多少万元？"),
            ("employee_count", "企业员工总数是多少人？"),
            ("rd_employee_count", "其中研发人员有多少人？"),
            ("patent_count", "企业拥有多少项专利？"),
            ("main_products", "主要产品有哪些？（请列出1-3个）"),
            ("market_share", "在细分市场的占有率大概是多少百分比？")
        ]
        
        logger.info("[专精特新] 信息处理器初始化完成")
    
    async def extract_from_documents(self, files: List) -> Dict[str, Any]:
        """从上传文档中提取企业信息"""
        extracted_data = {}
        
        for file_info in files:
            try:
                file_path = file_info.get("path", "")
                file_name = file_info.get("filename", "")
                
                if not os.path.exists(file_path):
                    logger.warning(f"[专精特新] 文件不存在: {file_path}")
                    continue
                
                # 根据文件类型调用不同的提取方法
                if file_name.lower().endswith('.pdf'):
                    pdf_data = await self._extract_from_pdf(file_path)
                    extracted_data.update(pdf_data)
                    
                elif file_name.lower().endswith(('.xlsx', '.xls')):
                    excel_data = await self._extract_from_excel(file_path)
                    extracted_data.update(excel_data)
                    
                elif file_name.lower().endswith(('.doc', '.docx')):
                    doc_data = await self._extract_from_word(file_path)
                    extracted_data.update(doc_data)
                    
                elif file_name.lower().endswith(('.jpg', '.jpeg', '.png')):
                    # 图片文件，可能是营业执照等
                    img_data = await self._extract_from_image(file_path, file_name)
                    extracted_data.update(img_data)
                    
            except Exception as e:
                logger.error(f"[专精特新] 提取文件 {file_info.get('filename')} 时出错: {e}")
        
        # 数据清洗和标准化
        cleaned_data = self._clean_extracted_data(extracted_data)
        
        logger.info(f"[专精特新] 从文档提取到 {len(cleaned_data)} 条信息")
        return cleaned_data
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """从PDF文件中提取信息"""
        data = {}
        
        try:
            # 尝试导入pdfplumber
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                text_content = ""
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
                
                # 从文本中提取关键信息
                data.update(self._extract_from_text(text_content))
                
                # 尝试提取表格数据
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        table_data = self._parse_table(table)
                        data.update(table_data)
                        
        except ImportError:
            logger.warning("[专精特新] pdfplumber未安装，使用简单文本提取")
            # 使用PyPDF2作为备选
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text_content = ""
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
                
                data.update(self._extract_from_text(text_content))
            except Exception as e:
                logger.error(f"[专精特新] PDF提取失败: {e}")
        
        return data
    
    async def _extract_from_excel(self, file_path: str) -> Dict[str, Any]:
        """从Excel文件中提取信息"""
        data = {}
        
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            # 检查常见的工作表名称
            sheet_names = wb.sheetnames
            target_sheets = ["财务数据", "审计报告", "财务", "数据", "Sheet1"]
            
            for sheet_name in sheet_names:
                if any(target in sheet_name for target in target_sheets):
                    ws = wb[sheet_name]
                    
                    # 提取所有单元格数据
                    cell_data = {}
                    for row in ws.iter_rows(values_only=True):
                        for cell in row:
                            if cell and isinstance(cell, (int, float, str)):
                                cell_str = str(cell).strip()
                                if cell_str and len(cell_str) > 1:
                                    # 尝试识别关键字段
                                    identified = self._identify_field(cell_str)
                                    if identified:
                                        cell_data.update(identified)
                    
                    data.update(cell_data)
                    
        except Exception as e:
            logger.error(f"[专精特新] Excel提取失败: {e}")
        
        return data
    
    async def _extract_from_word(self, file_path: str) -> Dict[str, Any]:
        """从Word文件中提取信息"""
        data = {}
        
        try:
            # 尝试使用python-docx
            import docx
            
            doc = docx.Document(file_path)
            text_content = ""
            
            # 提取所有段落文本
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # 提取表格数据
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    text_content += row_text + "\n"
            
            data.update(self._extract_from_text(text_content))
            
        except ImportError:
            logger.warning("[专精特新] python-docx未安装，使用纯文本提取")
            # 尝试以文本方式读取
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
                data.update(self._extract_from_text(text_content))
            except:
                pass
        
        return data
    
    async def _extract_from_image(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """从图片文件中提取信息（主要是识别文件类型）"""
        data = {}
        
        # 根据文件名判断
        if "营业执照" in file_name or "business" in file_name.lower():
            data["has_business_license"] = True
            data["business_license_file"] = file_name
            
        elif "审计报告" in file_name or "audit" in file_name.lower():
            data["has_audit_report"] = True
            data["audit_report_file"] = file_name
            
        elif "专利" in file_name or "patent" in file_name.lower():
            data["has_patent_cert"] = True
            data["patent_file"] = file_name
            
        elif "社保" in file_name or "social" in file_name.lower():
            data["has_social_security"] = True
            data["social_security_file"] = file_name
        
        return data
    
    def _extract_from_text(self, text: str) -> Dict[str, Any]:
        """从文本中提取关键信息"""
        data = {}
        
        # 企业名称提取
        company_patterns = [
            r"企业名称[：:]\s*([^\n]+)",
            r"公司名称[：:]\s*([^\n]+)",
            r"名称[：:]\s*([^\n]+)",
            r"^(.*?)(有限公司|有限责任公司|股份公司|集团公司)"
        ]
        
        for pattern in company_patterns:
            match = re.search(pattern, text)
            if match:
                company_name = match.group(1).strip()
                if len(company_name) > 2 and len(company_name) < 50:
                    data["company_name"] = company_name
                    break
        
        # 财务数据提取
        financial_patterns = {
            "revenue_2023": r"2023.*?营业收[入入][：:]\s*([\d,]+\.?\d*)\s*万?元?",
            "revenue_2022": r"2022.*?营业收[入入][：:]\s*([\d,]+\.?\d*)\s*万?元?",
            "revenue_2021": r"2021.*?营业收[入入][：:]\s*([\d,]+\.?\d*)\s*万?元?",
            "main_revenue_2023": r"2023.*?主营.*?收[入入][：:]\s*([\d,]+\.?\d*)\s*万?元?",
            "rd_expense_2023": r"2023.*?研发费[用用][：:]\s*([\d,]+\.?\d*)\s*万?元?",
            "employee_count": r"员工[总数]?[：:]\s*([\d,]+)\s*人?",
            "rd_employee_count": r"研发人员[：:]\s*([\d,]+)\s*人?",
            "patent_count": r"专利[：:]\s*([\d,]+)\s*项?",
        }
        
        for field, pattern in financial_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    value_str = match.group(1).replace(",", "")
                    if "." in value_str:
                        value = float(value_str)
                    else:
                        value = int(value_str)
                    data[field] = value
                except:
                    pass
        
        # 行业信息提取
        industry_keywords = ["智能制造", "生物医药", "电子信息", "新材料", "新能源", 
                           "节能环保", "高端装备", "人工智能", "集成电路", "软件"]
        
        for keyword in industry_keywords:
            if keyword in text:
                data["industry"] = keyword
                break
        
        # 主营业务提取（简单提取）
        business_section = self._extract_section(text, "主营业务", "经营范围")
        if business_section:
            data["main_business"] = business_section[:100]  # 截断前100字符
        
        return data
    
    def _extract_section(self, text: str, start_keyword: str, end_keyword: str = None) -> Optional[str]:
        """提取文本中的特定部分"""
        start_idx = text.find(start_keyword)
        if start_idx == -1:
            return None
        
        start_idx += len(start_keyword)
        
        if end_keyword:
            end_idx = text.find(end_keyword, start_idx)
            if end_idx == -1:
                section = text[start_idx:start_idx+500]
            else:
                section = text[start_idx:end_idx]
        else:
            # 提取到下一个标题或一定长度
            section = text[start_idx:start_idx+500]
        
        # 清理空白字符
        section = re.sub(r'\s+', ' ', section).strip()
        return section if section else None
    
    def _parse_table(self, table: List[List]) -> Dict[str, Any]:
        """解析表格数据"""
        data = {}
        
        try:
            # 将表格转换为键值对
            for row in table:
                if len(row) >= 2:
                    key = str(row[0]).strip() if row[0] else ""
                    value = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                    
                    if key and value:
                        identified = self._identify_field(f"{key}: {value}")
                        if identified:
                            data.update(identified)
        except:
            pass
        
        return data
    
    def _identify_field(self, text: str) -> Dict[str, Any]:
        """识别文本中的字段"""
        patterns = {
            "company_name": [r"企业名称", r"公司名称", r"名称"],
            "establishment_year": [r"成立时间", r"注册时间", r"成立"],
            "industry": [r"所属行业", r"行业"],
            "main_business": [r"主营业务", r"经营范围"],
            "revenue_2023": [r"2023.*收入", r"2023.*营收"],
            "rd_expense_2023": [r"2023.*研发", r"研发费用"],
            "employee_count": [r"员工总数", r"职工人数"],
            "patent_count": [r"专利数量", r"专利数"],
            "market_share": [r"市场占有率", r"市场份额"]
        }
        
        for field, keywords in patterns.items():
            for keyword in keywords:
                if keyword in text:
                    # 提取数值
                    numbers = re.findall(r"[\d,]+\.?\d*", text)
                    if numbers:
                        try:
                            value = float(numbers[0].replace(",", ""))
                            return {field: value}
                        except:
                            pass
                    # 提取文本
                    elif ":" in text or "：" in text:
                        parts = re.split(r"[：:]", text, 1)
                        if len(parts) > 1:
                            value = parts[1].strip()
                            if value and len(value) < 100:
                                return {field: value}
        
        return {}
    
    def _clean_extracted_data(self, data: Dict) -> Dict[str, Any]:
        """清洗和标准化提取的数据"""
        cleaned = {}
        
        for key, value in data.items():
            if value is None:
                continue
                
            if isinstance(value, str):
                value = value.strip()
                if not value:
                    continue
            
            # 类型转换和标准化
            if key.endswith(("_2021", "_2022", "_2023")):
                try:
                    if isinstance(value, str):
                        value = float(value.replace(",", ""))
                    cleaned[key] = float(value)
                except:
                    pass
            elif key in ["employee_count", "rd_employee_count", "patent_count"]:
                try:
                    if isinstance(value, str):
                        value = int(float(value.replace(",", "")))
                    cleaned[key] = int(value)
                except:
                    pass
            elif key == "market_share":
                try:
                    if isinstance(value, str):
                        # 提取百分比数值
                        match = re.search(r"(\d+\.?\d*)%", value)
                        if match:
                            cleaned[key] = float(match.group(1))
                        else:
                            numbers = re.findall(r"\d+\.?\d*", value)
                            if numbers:
                                cleaned[key] = float(numbers[0])
                    else:
                        cleaned[key] = float(value)
                except:
                    pass
            else:
                cleaned[key] = value
        
        return cleaned
    
    async def parse_text_input(self, text: str) -> Dict[str, Any]:
        """解析批量文本输入"""
        data = {}
        
        # 按行分割
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 尝试匹配"字段：值"格式
            if "：" in line or ":" in line:
                separator = "：" if "：" in line else ":"
                parts = line.split(separator, 1)
                if len(parts) == 2:
                    key = parts[0].strip()
                    value = parts[1].strip()
                    
                    # 映射字段名
                    mapped_key = self._map_field_name(key)
                    if mapped_key:
                        data[mapped_key] = value
        
        # 也尝试从整个文本中提取
        extracted = self._extract_from_text(text)
        data.update(extracted)
        
        return self._clean_extracted_data(data)
    
    def _map_field_name(self, field_name: str) -> Optional[str]:
        """映射中文字段名到标准字段名"""
        mapping = {
            "企业名称": "company_name",
            "公司名称": "company_name",
            "名称": "company_name",
            "成立时间": "establishment_year",
            "成立年份": "establishment_year",
            "注册时间": "establishment_year",
            "行业": "industry",
            "所属行业": "industry",
            "主营业务": "main_business",
            "主营收入": "main_revenue_2023",
            "营业收入": "revenue_2023",
            "营收": "revenue_2023",
            "研发费用": "rd_expense_2023",
            "研发投入": "rd_expense_2023",
            "员工总数": "employee_count",
            "职工人数": "employee_count",
            "研发人员": "rd_employee_count",
            "专利数量": "patent_count",
            "专利数": "patent_count",
            "主要产品": "main_products",
            "产品": "main_products",
            "市场占有率": "market_share",
            "市场份额": "market_share"
        }
        
        # 精确匹配
        if field_name in mapping:
            return mapping[field_name]
        
        # 模糊匹配
        for chinese_key, english_key in mapping.items():
            if chinese_key in field_name:
                return english_key
        
        return None
    
    def check_completeness(self, data: Dict[str, Any]) -> float:
        """检查数据完整性"""
        if not data:
            return 0.0
        
        required_count = 0
        optional_count = 0
        
        for field in self.REQUIRED_FIELDS:
            if field in data and data[field] not in [None, "", 0]:
                required_count += 1
        
        for field in self.OPTIONAL_FIELDS:
            if field in data and data[field] not in [None, "", 0]:
                optional_count += 1
        
        # 计算完整性百分比（必需字段权重70%，可选字段权重30%）
        required_score = (required_count / len(self.REQUIRED_FIELDS)) * 70
        optional_score = (optional_count / len(self.OPTIONAL_FIELDS)) * 30 if self.OPTIONAL_FIELDS else 0
        
        completeness = required_score + optional_score
        
        # 保存完整性到数据中
        data["completeness"] = completeness
        
        return completeness
    
    def get_missing_fields(self, data: Dict[str, Any]) -> List[str]:
        """获取缺失的必需字段"""
        missing = []
        
        for field in self.REQUIRED_FIELDS:
            if field not in data or data[field] in [None, "", 0]:
                missing.append(field)
        
        return missing
    
    def generate_questions(self, missing_fields: List[str], existing_data: Dict[str, Any]) -> List[str]:
        """根据缺失字段生成智能问题"""
        questions = []
        
        # 字段到问题的映射
        field_questions = {
            "company_name": "企业全称是什么？",
            "establishment_year": "企业成立时间是哪一年？",
            "industry": "企业所属行业是什么？（如：智能制造、生物医药、电子信息等）",
            "main_business": "主营业务具体是什么？",
            "revenue_2023": "2023年营业收入是多少万元？",
            "main_revenue_2023": "2023年主营业务收入是多少万元？",
            "rd_expense_2023": "2023年研发费用是多少万元？",
            "employee_count": "企业员工总数是多少人？",
            "rd_employee_count": "其中研发人员有多少人？",
            "patent_count": "企业拥有多少项专利？",
            "main_products": "主要产品有哪些？（请列出1-3个核心产品）",
            "market_share": "在细分市场的占有率大概是多少百分比？"
        }
        
        # 根据已有数据生成上下文相关的问题
        for field in missing_fields:
            if field in field_questions:
                question = field_questions[field]
                
                # 添加上下文信息
                if field == "market_share" and "industry" in existing_data:
                    question = f"在{existing_data['industry']}行业，{question}"
                elif field == "rd_expense_2023" and "revenue_2023" in existing_data:
                    question = f"2023年营收{existing_data['revenue_2023']}万元，{question}"
                
                questions.append(question)
        
        return questions
    
    async def handle_conversation(self, user_input: str, current_step: int, existing_data: Dict[str, Any]) -> Dict[str, Any]:
        """处理对话引导"""
        result = {
            "extracted_data": {},
            "next_step": current_step,
            "completed": False,
            "next_question": ""
        }
        
        if current_step >= len(self.conversation_steps):
            result["completed"] = True
            return result
        
        # 获取当前步骤的字段和问题
        current_field, current_question = self.conversation_steps[current_step]
        
        # 从用户输入中提取信息
        extracted = self._extract_field_from_response(current_field, user_input)
        if extracted:
            result["extracted_data"][current_field] = extracted
        
        # 更新步骤
        result["next_step"] = current_step + 1
        
        # 检查是否完成
        if result["next_step"] >= len(self.conversation_steps):
            result["completed"] = True
            result["next_question"] = "信息收集完成！"
        else:
            # 获取下一个问题
            next_field, next_question = self.conversation_steps[result["next_step"]]
            result["next_question"] = next_question
        
        return result
    
    def _extract_field_from_response(self, field: str, response: str) -> Any:
        """从用户响应中提取特定字段信息"""
        if field == "company_name":
            # 提取企业名称
            patterns = [
                r"^(.*?)(有限公司|有限责任公司|股份公司|集团公司)",
                r"企业[名称名]?[：:]\s*(.+)",
                r"公司[名称名]?[：:]\s*(.+)"
            ]
            
            for pattern in patterns:
                match = re.search(pattern, response)
                if match:
                    name = match.group(1).strip() if match.group(1) else match.group(0).strip()
                    if len(name) > 1:
                        return name
        
        elif field in ["establishment_year", "revenue_2023", "main_revenue_2023", 
                      "rd_expense_2023", "employee_count", "rd_employee_count", 
                      "patent_count", "market_share"]:
            # 提取数值
            numbers = re.findall(r"[\d,]+\.?\d*", response)
            if numbers:
                try:
                    value = float(numbers[0].replace(",", ""))
                    
                    # 特殊处理：成立年份应该是整数
                    if field == "establishment_year":
                        return int(value)
                    
                    # 市场占有率可能是百分比
                    if field == "market_share" and "%" in response:
                        return value
                    elif field == "market_share":
                        # 如果没有百分比符号，假设是百分比值
                        return value
                    
                    return value
                except:
                    pass
        
        elif field in ["industry", "main_business", "main_products"]:
            # 提取文本信息
            # 移除常见的问题前缀
            response_clean = re.sub(r"^(企业|公司|我[们们]的?)?(所属)?(行业|业务|产品)[是是为]?", "", response).strip()
            if response_clean and len(response_clean) > 1:
                return response_clean[:200]  # 限制长度
        
        # 如果无法提取，返回原始响应（后续可以进一步处理）
        return response.strip() if response.strip() else None
    
    def validate_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """验证数据的合理性"""
        validated = data.copy()
        
        # 验证数值范围
        if "establishment_year" in validated:
            year = validated["establishment_year"]
            if isinstance(year, (int, float)):
                if year < 1900 or year > 2024:
                    validated["establishment_year"] = None
        
        if "revenue_2023" in validated:
            revenue = validated["revenue_2023"]
            if isinstance(revenue, (int, float)):
                if revenue < 0 or revenue > 10000000:  # 假设不超过100亿
                    validated["revenue_2023"] = None
        
        if "market_share" in validated:
            share = validated["market_share"]
            if isinstance(share, (int, float)):
                if share < 0 or share > 100:
                    validated["market_share"] = None
        
        # 计算衍生指标
        self._calculate_derived_metrics(validated)
        
        return validated
    
    def _calculate_derived_metrics(self, data: Dict[str, Any]):
        """计算衍生指标"""
        # 主营业务占比
        if "main_revenue_2023" in data and "revenue_2023" in data:
            main_revenue = data["main_revenue_2023"]
            total_revenue = data["revenue_2023"]
            if total_revenue and total_revenue > 0:
                data["main_revenue_ratio"] = (main_revenue / total_revenue) * 100
        
        # 研发费用占比
        if "rd_expense_2023" in data and "revenue_2023" in data:
            rd_expense = data["rd_expense_2023"]
            total_revenue = data["revenue_2023"]
            if total_revenue and total_revenue > 0:
                data["rd_expense_ratio"] = (rd_expense / total_revenue) * 100
        
        # 研发人员占比
        if "rd_employee_count" in data and "employee_count" in data:
            rd_employees = data["rd_employee_count"]
            total_employees = data["employee_count"]
            if total_employees and total_employees > 0:
                data["rd_employee_ratio"] = (rd_employees / total_employees) * 100
        
        # 企业年龄
        if "establishment_year" in data:
            establishment_year = data["establishment_year"]
            if establishment_year:
                current_year = 2024  # 假设当前年份
                data["company_age"] = current_year - establishment_year
    
    def get_data_summary(self, data: Dict[str, Any]) -> str:
        """获取数据摘要"""
        summary = "【企业信息摘要】\n\n"
        
        fields_to_show = [
            ("company_name", "企业名称"),
            ("establishment_year", "成立时间"),
            ("industry", "所属行业"),
            ("main_business", "主营业务"),
            ("revenue_2023", "2023年营收（万元）"),
            ("main_revenue_ratio", "主营业务占比（%）"),
            ("rd_expense_ratio", "研发费用占比（%）"),
            ("employee_count", "员工总数"),
            ("rd_employee_ratio", "研发人员占比（%）"),
            ("patent_count", "专利数量"),
            ("market_share", "市场占有率（%）")
        ]
        
        for field_en, field_cn in fields_to_show:
            if field_en in data and data[field_en] not in [None, "", 0]:
                value = data[field_en]
                if isinstance(value, float):
                    value = f"{value:.2f}"
                summary += f"{field_cn}: {value}\n"
        
        completeness = self.check_completeness(data)
        summary += f"\n信息完整性: {completeness:.1f}%"
        
        return summary
